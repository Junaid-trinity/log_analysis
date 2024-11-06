import streamlit as st
import asyncio
from data_ingestion import DataIngestion
from data_processing import DataProcessing_csv
from agents import ImageerrorAnalysisagent , LogFilesAnalysisagentCSV , LogFilesAnalysisAgentLog 
from dotenv import load_dotenv
from io import BytesIO
from PIL import Image
import base64
import os
import tempfile
import json
from docx import Document
from docx.shared import Pt
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


load_dotenv(override=True)

async def main():
    # Initialize classes
    data_ingestion = DataIngestion()
    data_processing = DataProcessing_csv()
    Image_error_Analysis_agent = ImageerrorAnalysisagent()
    LogFilesAnalysisagent_csv = LogFilesAnalysisagentCSV()
    LogFilesAnalysisagent_log = LogFilesAnalysisAgentLog()

    # Set Streamlit layout options
    st.set_page_config(page_title="GenAI Error Log Analysis", layout="wide")
    # Sidebar Configuration
    with st.sidebar:
        st.title("Upload Options")
        gitlab_logs = None
        log_files_gitlab = None
        if "log_files_gitlab" not in st.session_state:
            st.session_state.log_files_gitlab = None
        # Fetch projects asynchronously
        # File uploader for error log
        uploaded_files = st.file_uploader("Upload your error log (csv,log) or error screenshots (png,jpg,jpeg)", type=["csv", "log","png", "jpg", "jpeg"], accept_multiple_files=True)

        
        # Display "OR" separator
        st.markdown("<style>.separator {display: flex; align-items: center; text-align: center; color: #888; margin: 20px 0;} .separator::before, .separator::after {content: ''; flex: 1; border-bottom: 1px solid #ccc;} .separator::before {margin-right: 10px;} .separator::after {margin-left: 10px;}</style><div class='separator'>OR</div>", unsafe_allow_html=True)

        projects = await data_ingestion.get_gitlab_projects()
        projects = [(None, None)] + projects
        # Add a project selection dropdown
        selected_project_id = st.selectbox("Select a Project from gitlab", options=projects, format_func=lambda x: x[1] if x[1] else None)
        
        
        if  selected_project_id[0] is not None:
            st.session_state.selected_project_id = selected_project_id[0]  # Store only the ID
            branches = await data_ingestion.get_gitlab_branches(st.session_state.selected_project_id)
            branches = [(None, "Select a Branch")] + branches
            # Display "OR" separator
            st.markdown("<style>.separator {display: flex; align-items: center; text-align: center; color: #888; margin: 20px 0;} .separator::before, .separator::after {content: ''; flex: 1; border-bottom: 1px solid #ccc;} .separator::before {margin-right: 10px;} .separator::after {margin-left: 10px;}</style><div class='separator'>AND</div>", unsafe_allow_html=True)
            selected_branch = st.selectbox("Select a Branch", options=branches, format_func=lambda x: x[1])
            
            if selected_branch and selected_branch[0] is not None:
                st.session_state.selected_branch_name = selected_branch[1]
            
             # Step 3: Ask for the log file path from the user
            # Display "OR" separator
            st.markdown("<style>.separator {display: flex; align-items: center; text-align: center; color: #888; margin: 20px 0;} .separator::before, .separator::after {content: ''; flex: 1; border-bottom: 1px solid #ccc;} .separator::before {margin-right: 10px;} .separator::after {margin-left: 10px;}</style><div class='separator'>AND</div>", unsafe_allow_html=True)
            log_path = st.text_input("Enter the path to log files (e.g., `/` for root, `logs` for logs folder, `logs/errors` for nested folders)")

            gitlab_logs =  st.button("Fetch Log Files from gitlab")
            
            
            
        
        # Display "OR" separator
#        st.markdown("<style>.separator {display: flex; align-items: center; text-align: center; color: #888; margin: 20px 0;} .separator::before, .separator::after {content: ''; flex: 1; border-bottom: 1px solid #ccc;} .separator::before {margin-right: 10px;} .separator::after {margin-left: 10px;}</style><div class='separator'>OR</div>", unsafe_allow_html=True)
        
        # Screenshot uploader for multiple files
#        screenshot_files = st.file_uploader("Upload screenshots (PNG or JPG)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    
    # Main Window
    st.title("Automated Error Log Analysis")
    st.write(
        "Use this tool to automate the analysis of error logs with GenAI. "
        "Upload log files or error screenshots to get detailed error reports."
    )
    
    # Header row with Analyze and Download buttons
    col1, col2 = st.columns([3, 1])
    with col1:
        analyze_button = st.button("Analyze Error Log")
    with col2:
        # Placeholder for the download button to appear later
        download_button_placeholder = st.empty()
    
    if gitlab_logs:
        with st.spinner("Fetching log files..."):
            st.session_state.log_files_gitlab = await data_ingestion.get_gitlab_log_files(st.session_state.selected_project_id,st.session_state.selected_branch_name,log_path)
#            st.write(log_files)
            if st.session_state.log_files_gitlab:
                for log_file in st.session_state.log_files_gitlab:
                    st.write(log_file.get('name'))  # 
    
    
    if analyze_button:
        with st.spinner("Processing and analyzing logs... This may take a moment."):
            if uploaded_files:
                csv_files = [file for file in uploaded_files if file.name.endswith(".csv")]
                log_files = [file for file in uploaded_files if file.name.endswith(".log")]
                screenshot_files = [file for file in uploaded_files if file.name.endswith((".png", ".jpg", ".jpeg"))]
                
                if csv_files:
                    data_frames = await data_ingestion.load_error_log(uploaded_files)
                    log_file_responses = await LogFilesAnalysisagent_csv.execute_agent(data_frames)
                
                    # Display each result
                    for i, data in enumerate(log_file_responses):
                        for batch_report in data:  # Each 'data' item is a list of batch reports
                            st.write(batch_report)
                
                if log_files:
                    temp_log_files = {}
                    file_number = 1  # File counter
                    for log_file in log_files:
                        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                            temp_file.write(log_file.read())
                            temp_log_files[temp_file.name] = log_file.name
                    
                    temp_file_paths = list(temp_log_files.keys())
                    analysis_responses  = await LogFilesAnalysisagent_log.execute_agent(temp_file_paths)
                    doc = Document()
                    heading = doc.add_heading("Error Log Analysis Report", level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the heading
                    st.markdown("""
                            <style>
                                .custom-box {
                                    background-color: #f0f4f8;  /* Light grey-blue background */
                                    padding: 15px;
                                    border-radius: 8px;
                                    margin-bottom: 10px;
                                }
                                .custom-heading {
                                    color: #333333;  /* Darker text color for better readability */
                                    font-weight: bold;
                                    font-size: 24px;  /* Increase font size for the heading */
                                    margin-bottom: 10px;
                                }
                            </style>
                        """, unsafe_allow_html=True)
                    for  result in analysis_responses:
                        response = result[0]
                        original_file_name = temp_log_files[result[1]]
                        # Display the file name as a subheader
                        st.subheader(f"**File {file_number}:** {original_file_name}")
                        doc.add_heading(f"File {file_number}: {original_file_name}", level=2)
                        file_number += 1

                        # Parse and display the detailed response
                        error_analyzed = json.loads(response).get('DetailedResponseFull', [])
                        st.write("### Error Report")
                        error_number = 1  # Error counter within the file
                        for error in error_analyzed:
                            st.markdown(f"<div class='custom-box'><div class='custom-heading'> {error_number}: {error.get('heading', 'N/A')}</div></div>", unsafe_allow_html=True)
                            doc.add_paragraph(f"{error_number}: {error.get('heading', 'N/A')}", style="Heading2")
                            error_number += 1
                            if error.get('Date'):
                                st.write(f"**Date:** {error.get('Date')}")
                                doc.add_paragraph(f"**Date:** {error.get('Date')}")
                            st.write(f"**Error Message:** {error.get('error_message', 'N/A')}")
                            doc.add_paragraph(f"**Error Message:** {error.get('error_message', 'N/A')}")
                            st.write(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in logs'}")
                            doc.add_paragraph(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in logs'}")
                            st.write(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in logs'}")
                            doc.add_paragraph(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in logs'}")
                            st.write(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
                            doc.add_paragraph(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
        #                     st.write(f"**Suggested Solutions:** {error.get('suggested_solutions', 'N/A')}")
        #                     st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                            suggested_solution = error.get('suggested_solutions', {})
                            st.write(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
                            doc.add_paragraph(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
                
                            # Display Original and Corrected Code only if they exist
                            code_fix = suggested_solution.get('code_fix')
                            if code_fix:
                                    original_code = code_fix.get('original_code', None)
                                    corrected_code = code_fix.get('corrected_code', None)
                                    
                                    if original_code:
                                        st.write("**Original Code:**")
                                        doc.add_paragraph("Original Code:", style="Heading3")
                                        st.code(original_code, language='python')
                                        # Add original code block with monospaced font
                                        original_code_paragraph = doc.add_paragraph()
                                        original_code_run = original_code_paragraph.add_run(original_code)
                                        original_code_run.font.name = 'Courier New'
                                        original_code_run.font.size = Pt(10)
                                        
                                        
                                        # Optional: Add a border around the code block
                                        original_code_paragraph_format = original_code_paragraph.paragraph_format
                                        p = original_code_paragraph._element
                                        pPr = p.get_or_add_pPr()
                                        pBdr = OxmlElement("w:pBdr")
                                        for border_name in ["top", "bottom", "left", "right"]:
                                            border = OxmlElement(f"w:{border_name}")
                                            border.set(qn("w:val"), "single")
                                            border.set(qn("w:sz"), "4")
                                            border.set(qn("w:space"), "1")
                                            border.set(qn("w:color"), "auto")
                                            pBdr.append(border)
                                        pPr.append(pBdr)
                                        
                                        
                                        
                                    if corrected_code:
                                        st.write("**Corrected Code:**")
                                        doc.add_paragraph("Corrected Code:", style="Heading3")
                                        st.code(corrected_code, language='python')
                                        # Add corrected code block with monospaced font
                                        corrected_code_paragraph = doc.add_paragraph()
                                        corrected_code_run = corrected_code_paragraph.add_run(corrected_code)
                                        corrected_code_run.font.name = 'Courier New'
                                        corrected_code_run.font.size = Pt(10)
                                        
                                        
                                                                    # Optional: Add a border around the corrected code block
                                        corrected_code_paragraph_format = corrected_code_paragraph.paragraph_format
                                        p = corrected_code_paragraph._element
                                        pPr = p.get_or_add_pPr()
                                        pBdr = OxmlElement("w:pBdr")
                                        for border_name in ["top", "bottom", "left", "right"]:
                                            border = OxmlElement(f"w:{border_name}")
                                            border.set(qn("w:val"), "single")
                                            border.set(qn("w:sz"), "4")
                                            border.set(qn("w:space"), "1")
                                            border.set(qn("w:color"), "auto")
                                            pBdr.append(border)
                                        pPr.append(pBdr)
                                
                            st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                            doc.add_paragraph(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                            doc.add_paragraph("\n")
                            st.markdown('</div>', unsafe_allow_html=True)  # Close custom div  
                            
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)       
                    download_button_placeholder.download_button(
                        label="Download Error Report",
                        data=buffer,
                        file_name="error_log_analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )       
                            
                if screenshot_files:
                    file_number = 1  # File counter
                    base_64_files = await data_processing.image_to_base64(screenshot_files)
                    # Display each image using its base64 representation
                    image_responses = await Image_error_Analysis_agent.execute_agent(base_64_files)
                    doc = Document()
                    heading = doc.add_heading("Error Log Analysis Report", level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the heading
                    for file_name, response in zip(screenshot_files, image_responses):
                        st.subheader(f"**File {file_number}:** {file_name.name}")  # Display the file name
                        doc.add_heading(f"File {file_number}: {file_name.name}", level=2)
                        file_number += 1
                        
                        error_anlaysed = json.loads(response).get('DetailedResponseFull')
                        st.write("### Error Report")
                        st.markdown("""
                <style>
                    .custom-box {
                        background-color: #f0f4f8;  /* Light grey-blue background */
                        padding: 15px;
                        border-radius: 8px;
                        margin-bottom: 10px;
                    }
                    .custom-heading {
                        color: #333333;  /* Darker text color for better readability */
                        font-weight: bold;
                        font-size: 24px;  /* Increase font size for the heading */
                        margin-bottom: 10px;
                    }
                </style>
            """, unsafe_allow_html=True)
                        error_number = 1  # Error counter within the file
                        for error in error_anlaysed:
                # Heading inside the custom box
    #                        st.markdown(f"<div class='custom-box'><div class='custom-heading'>{error.get('heading', 'N/A')}</div></div>", unsafe_allow_html=True)
                            st.markdown(f"<div class='custom-box'><div class='custom-heading'> {error_number}: {error.get('heading', 'N/A')}</div></div>", unsafe_allow_html=True)
                            doc.add_paragraph(f"{error_number}: {error.get('heading', 'N/A')}", style="Heading2")
                            error_number += 1
                            if error.get('Date'):
                                st.write(f"**Date:** {error.get('Date')}")
                                doc.add_paragraph(f"**Date:** {error.get('Date')}")
                            st.write(f"**Error Message:** {error.get('error_message', 'N/A')}")
                            doc.add_paragraph(f"**Error Message:** {error.get('error_message', 'N/A')}")
                            st.write(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in screenshot'}")
                            doc.add_paragraph(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in screenshot'}")
                            st.write(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in screenshot'}")
                            doc.add_paragraph(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in screenshot'}")
                            st.write(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
                            doc.add_paragraph(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
                            
                            suggested_solution = error.get('suggested_solutions', {})
                            st.write(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
                            doc.add_paragraph(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
                            
                            # Display Original and Corrected Code only if they exist
                            code_fix = suggested_solution.get('code_fix')
                            if code_fix:
                                original_code = code_fix.get('original_code', None)
                                corrected_code = code_fix.get('corrected_code', None)
                                
                                if original_code:
                                    st.write("**Original Code:**")
                                    doc.add_paragraph("Original Code:", style="Heading3")
                                    st.code(original_code, language='python')
                                    # Add original code block with monospaced font
                                    original_code_paragraph = doc.add_paragraph()
                                    original_code_run = original_code_paragraph.add_run(original_code)
                                    original_code_run.font.name = 'Courier New'
                                    original_code_run.font.size = Pt(10)
                                    
                                    
                                    # Optional: Add a border around the code block
                                    original_code_paragraph_format = original_code_paragraph.paragraph_format
                                    p = original_code_paragraph._element
                                    pPr = p.get_or_add_pPr()
                                    pBdr = OxmlElement("w:pBdr")
                                    for border_name in ["top", "bottom", "left", "right"]:
                                        border = OxmlElement(f"w:{border_name}")
                                        border.set(qn("w:val"), "single")
                                        border.set(qn("w:sz"), "4")
                                        border.set(qn("w:space"), "1")
                                        border.set(qn("w:color"), "auto")
                                        pBdr.append(border)
                                    pPr.append(pBdr)
                                                                
                                    
                                    
                                
                                if corrected_code:
                                    st.write("**Corrected Code:**")
                                    doc.add_paragraph("Corrected Code:", style="Heading3")
                                    st.code(corrected_code, language='python')
                                    # Add corrected code block with monospaced font
                                    corrected_code_paragraph = doc.add_paragraph()
                                    corrected_code_run = corrected_code_paragraph.add_run(corrected_code)
                                    corrected_code_run.font.name = 'Courier New'
                                    corrected_code_run.font.size = Pt(10)
                                    
                                    
                                                                # Optional: Add a border around the corrected code block
                                    corrected_code_paragraph_format = corrected_code_paragraph.paragraph_format
                                    p = corrected_code_paragraph._element
                                    pPr = p.get_or_add_pPr()
                                    pBdr = OxmlElement("w:pBdr")
                                    for border_name in ["top", "bottom", "left", "right"]:
                                        border = OxmlElement(f"w:{border_name}")
                                        border.set(qn("w:val"), "single")
                                        border.set(qn("w:sz"), "4")
                                        border.set(qn("w:space"), "1")
                                        border.set(qn("w:color"), "auto")
                                        pBdr.append(border)
                                    pPr.append(pBdr)
                                    
                                    
                            
                            st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                            doc.add_paragraph(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                            doc.add_paragraph("\n")
                            st.markdown('</div>', unsafe_allow_html=True)  # Close custom div           
                    
                            
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)       
                    download_button_placeholder.download_button(
                        label="Download Error Report",
                        data=buffer,
                        file_name="error_log_analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )        
                
            if st.session_state.log_files_gitlab:
                # Temporarily store GitLab files in a dictionary similar to uploaded log files
                temp_log_files_gitlab = {}
                file_number = 1  # File counter
                for gitlab_file in st.session_state.log_files_gitlab:
                    file_content = gitlab_file.get("content")
                    if file_content is not None:  # Check if content is not None
                        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                            temp_file.write(file_content)  # Write content to temp file
                            temp_log_files_gitlab[temp_file.name] = gitlab_file.get("name")
                    else:
                        # Optional: log or handle cases where content is None
                        print(f"Warning: No content found for file {gitlab_file.get('name')}")
                temp_file_paths_gitlab = list(temp_log_files_gitlab.keys())
                analysis_responses_gitlab = await LogFilesAnalysisagent_log.execute_agent(temp_file_paths_gitlab)
                doc = Document()
                heading = doc.add_heading("Error Log Analysis Report", level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the heading
                st.markdown("""
                        <style>
                            .custom-box {
                                background-color: #f0f4f8;  /* Light grey-blue background */
                                padding: 15px;
                                border-radius: 8px;
                                margin-bottom: 10px;
                            }
                            .custom-heading {
                                color: #333333;  /* Darker text color for better readability */
                                font-weight: bold;
                                font-size: 24px;  /* Increase font size for the heading */
                                margin-bottom: 10px;
                            }
                        </style>
                    """, unsafe_allow_html=True)
                for  result in analysis_responses_gitlab:
                    response = result[0]
                    original_file_name = temp_log_files_gitlab[result[1]]
                    # Display the file name as a subheader
                    st.subheader(f"**File {file_number}:** {original_file_name}")
                    doc.add_heading(f"File {file_number}: {original_file_name}", level=2)
                    file_number += 1

                    # Parse and display the detailed response
                    error_analyzed = json.loads(response).get('DetailedResponseFull', [])
                    st.write("### Error Report")
                    error_number = 1  # Error counter within the file
                    for error in error_analyzed:
                        st.markdown(f"<div class='custom-box'><div class='custom-heading'> {error_number}: {error.get('heading', 'N/A')}</div></div>", unsafe_allow_html=True)
                        doc.add_paragraph(f"{error_number}: {error.get('heading', 'N/A')}", style="Heading2")
                        error_number += 1
                        if error.get('Date'):
                                st.write(f"**Date:** {error.get('Date')}")
                                doc.add_paragraph(f"**Date:** {error.get('Date')}")
                        st.write(f"**Error Message:** {error.get('error_message', 'N/A')}")
                        doc.add_paragraph(f"**Error Message:** {error.get('error_message', 'N/A')}")
                        st.write(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in logs'}")
                        doc.add_paragraph(f"**Line Number:** {error.get('line_number') if error.get('line_number') is not None else 'Information not available in logs'}")
                        st.write(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in logs'}")
                        doc.add_paragraph(f"**File Impacted:** {error.get('file_impacted') if error.get('file_impacted') is not None else 'Information not available in logs'}")
                        st.write(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
                        doc.add_paragraph(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
    #                     st.write(f"**Suggested Solutions:** {error.get('suggested_solutions', 'N/A')}")
    #                     st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                        suggested_solution = error.get('suggested_solutions', {})
                        st.write(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
                        doc.add_paragraph(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
            
                        # Display Original and Corrected Code only if they exist
                        code_fix = suggested_solution.get('code_fix')
                        if code_fix:
                                original_code = code_fix.get('original_code', None)
                                corrected_code = code_fix.get('corrected_code', None)
                                
                                if original_code:
                                    st.write("**Original Code:**")
                                    doc.add_paragraph("Original Code:", style="Heading3")
                                    st.code(original_code, language='python')
                                    # Add original code block with monospaced font
                                    original_code_paragraph = doc.add_paragraph()
                                    original_code_run = original_code_paragraph.add_run(original_code)
                                    original_code_run.font.name = 'Courier New'
                                    original_code_run.font.size = Pt(10)
                                    
                                    
                                    # Optional: Add a border around the code block
                                    original_code_paragraph_format = original_code_paragraph.paragraph_format
                                    p = original_code_paragraph._element
                                    pPr = p.get_or_add_pPr()
                                    pBdr = OxmlElement("w:pBdr")
                                    for border_name in ["top", "bottom", "left", "right"]:
                                        border = OxmlElement(f"w:{border_name}")
                                        border.set(qn("w:val"), "single")
                                        border.set(qn("w:sz"), "4")
                                        border.set(qn("w:space"), "1")
                                        border.set(qn("w:color"), "auto")
                                        pBdr.append(border)
                                    pPr.append(pBdr)
                                    
                                    
                                    
                                if corrected_code:
                                    st.write("**Corrected Code:**")
                                    doc.add_paragraph("Corrected Code:", style="Heading3")
                                    st.code(corrected_code, language='python')
                                    # Add corrected code block with monospaced font
                                    corrected_code_paragraph = doc.add_paragraph()
                                    corrected_code_run = corrected_code_paragraph.add_run(corrected_code)
                                    corrected_code_run.font.name = 'Courier New'
                                    corrected_code_run.font.size = Pt(10)
                                    
                                    
                                                                # Optional: Add a border around the corrected code block
                                    corrected_code_paragraph_format = corrected_code_paragraph.paragraph_format
                                    p = corrected_code_paragraph._element
                                    pPr = p.get_or_add_pPr()
                                    pBdr = OxmlElement("w:pBdr")
                                    for border_name in ["top", "bottom", "left", "right"]:
                                        border = OxmlElement(f"w:{border_name}")
                                        border.set(qn("w:val"), "single")
                                        border.set(qn("w:sz"), "4")
                                        border.set(qn("w:space"), "1")
                                        border.set(qn("w:color"), "auto")
                                        pBdr.append(border)
                                    pPr.append(pBdr)
                            
                        st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                        doc.add_paragraph(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                        doc.add_paragraph("\n")
                        st.markdown('</div>', unsafe_allow_html=True)  # Close custom div  
                        
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)       
                download_button_placeholder.download_button(
                    label="Download Error Report",
                    data=buffer,
                    file_name="error_log_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        
    # Check if a project is selected and fetch log files
#    if  selected_project_id[0] is not None:
#        st.session_state.selected_project_id = selected_project_id[0]  # Store only the ID
        
        # Fetch log files only if a valid project is selected
#        log_file_paths = await data_ingestion.fetch_log_files(st.session_state.selected_project_id)
        
#        if log_file_paths:
#            for log_file in log_file_paths:
#                st.write(log_file)  # Display the full path of each log file
#        else:
#            st.write("No log files found in the selected project.")
    
    
                    
                
                
                
                
#                with st.chat_message("user"):
#                    st.subheader(error.get('heading', 'N/A'))
#                    st.write(f"**Error Message:** {error.get('error_message', 'N/A')}")
#                    st.write(f"**Line Number:** {error.get('line_number', 'N/A')}")
#                    st.write(f"**File Impacted:** {error.get('file_impacted', 'N/A')}")
#                    st.write(f"**Possible Cause(s):** {error.get('possible_cause', 'N/A')}")
#    #                     st.write(f"**Suggested Solutions:** {error.get('suggested_solutions', 'N/A')}")
#    #                     st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
#                    suggested_solution = error.get('suggested_solutions', {})
#                    st.write(f"**Suggested Solution:** {suggested_solution.get('suggested_solution', 'N/A')}")
            
                        # Display Original and Corrected Code only if they exist
#                    code_fix = suggested_solution.get('code_fix')
#                    if code_fix:
#                        original_code = code_fix.get('original_code', None)
#                        corrected_code = code_fix.get('corrected_code', None)
                                
#                        if original_code:
#                                    st.write("**Original Code:**")
#                                    st.code(original_code, language='python')
                                    
#                        if corrected_code:
#                                    st.write("**Corrected Code:**")
#                                    st.code(corrected_code, language='python')
                            
#                    st.write(f"**Associated Merge Requests:** {error.get('associated_merge_requests', 'N/A')}")
                
                    
            
    
        
        
# Run the Streamlit app
if __name__ == "__main__":
    asyncio.run(main())